"""
Serveur Flask pour remplir automatiquement le fichier SUIVI DE FORMATION
Ce serveur reçoit les données de Make.com et génère le document Word rempli
"""

from flask import Flask, request, jsonify, send_file
from docx import Document
import os
import json
from datetime import datetime
import tempfile

app = Flask(__name__)

# Configuration
TEMPLATE_PATH = os.environ.get('TEMPLATE_PATH', '/templates/SUIVI_DE_FORMATION_VIERGE.docx')

def fill_suivi_formation(data, template_path):
    """
    Remplit le document SUIVI DE FORMATION avec les données fournies
    
    Args:
        data: Dictionnaire contenant les données à remplir
        template_path: Chemin vers le fichier template
    
    Returns:
        Chemin vers le fichier généré
    """
    # Charger le template
    doc = Document(template_path)
    
    # TABLE 0: Informations générales (9 lignes x 2 colonnes)
    table0 = doc.tables[0]
    
    # Ligne 0-1: AFFECTATION
    if 'affectation' in data:
        table0.rows[0].cells[1].text = data['affectation']
        table0.rows[1].cells[1].text = data['affectation']
    
    # Ligne 2: Semaine
    if 'semaine' in data:
        table0.rows[2].cells[1].text = data['semaine']
    
    # Ligne 3: Nom du Formateur
    if 'formateur' in data:
        table0.rows[3].cells[1].text = data['formateur']
    
    # Ligne 4: Nom du référent
    if 'referent' in data:
        table0.rows[4].cells[1].text = data['referent']
    
    # Ligne 5: Horaires
    if 'horaires' in data:
        table0.rows[5].cells[1].text = data['horaires']
    
    # Ligne 6: Numéro d'action
    if 'numero_action' in data:
        table0.rows[6].cells[1].text = data['numero_action']
    
    # Ligne 7: Date de rédaction
    if 'date_redaction' in data:
        table0.rows[7].cells[1].text = data['date_redaction']
    else:
        table0.rows[7].cells[1].text = datetime.now().strftime('%d/%m/%Y')
    
    # Ligne 8: Observations sur le groupe
    if 'observations_groupe' in data:
        table0.rows[8].cells[1].text = data['observations_groupe']
    
    # TABLE 2: Thèmes et modules (2 lignes x 2 colonnes)
    table2 = doc.tables[2]
    
    # Ligne 0: THEMES ET MODULES ABORDES
    if 'themes_modules' in data:
        table2.rows[0].cells[1].text = data['themes_modules']
    
    # Ligne 1: PREVISION PROCHAINE SESSION
    if 'previsions' in data:
        table2.rows[1].cells[1].text = data['previsions']
    
    # TABLE 3: Apprenants (10 lignes x 3 colonnes)
    # Ligne 0 = en-têtes, lignes 1-9 = apprenants
    table3 = doc.tables[3]
    
    if 'apprenants' in data:
        apprenants = data['apprenants']
        for i, apprenant in enumerate(apprenants[:9]):  # Maximum 9 apprenants
            row_index = i + 1  # +1 car ligne 0 = en-têtes
            
            # Colonne 0: Nom
            if 'nom' in apprenant:
                table3.rows[row_index].cells[0].text = apprenant['nom']
            
            # Colonne 1: Prénom
            if 'prenom' in apprenant:
                table3.rows[row_index].cells[1].text = apprenant['prenom']
            
            # Colonne 2: Observation
            if 'observation' in apprenant:
                table3.rows[row_index].cells[2].text = apprenant['observation']
    
    # Générer le nom du fichier de sortie
    # Nettoyer l'affectation et la semaine pour créer un nom de fichier valide
    affectation = data.get('affectation', 'FORMATION').replace('/', '-').replace(' ', '_')
    semaine = data.get('semaine', '').replace('Du ', '').replace(' au ', '_au_').replace('/', '-').replace(' ', '_').strip()
    output_filename = f"SUIVI_FORMATION_{affectation}_{semaine}.docx"
    
    # Créer un fichier temporaire
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, output_filename)
    
    # Sauvegarder le document
    doc.save(output_path)
    
    return output_path, output_filename


@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint de vérification de santé du serveur"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    })


@app.route('/fill-document', methods=['POST'])
def fill_document():
    """
    Endpoint principal pour remplir le document
    
    Exemple de payload JSON:
    {
        "affectation": "CAP 2 OLBER",
        "semaine": "Du 24/03/2025 au 26/03/2025",
        "formateur": "Yves Sournac",
        "referent": "Jean-François SOLLEAU",
        "horaires": "8h-12h / 13h-17h",
        "numero_action": "25SF1353",
        "date_redaction": "26/03/2025",
        "observations_groupe": "Bonne dynamique de groupe",
        "themes_modules": "CP1 - Réception de marchandises\\nContrôle quantitatif",
        "previsions": "CP2 - Stockage et rangement",
        "apprenants": [
            {
                "nom": "DUPONT",
                "prenom": "Jean",
                "observation": "Très investi, bonne compréhension"
            },
            {
                "nom": "MARTIN",
                "prenom": "Sophie",
                "observation": "Quelques difficultés sur le scan"
            }
        ]
    }
    """
    try:
        # Récupérer les données JSON
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'Aucune donnée fournie'}), 400
        
        # Remplir le document
        output_path, output_filename = fill_suivi_formation(data, TEMPLATE_PATH)
        
        # Retourner le fichier
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({
            'error': str(e),
            'type': type(e).__name__
        }), 500


@app.route('/test-fill', methods=['GET'])
def test_fill():
    """Endpoint de test avec des données d'exemple"""
    test_data = {
        "affectation": "CAP 2 OLBER - TEST",
        "semaine": "Du 24/03/2025 au 26/03/2025",
        "formateur": "Yves Sournac",
        "referent": "Jean-François SOLLEAU",
        "horaires": "8h-12h / 13h-17h",
        "numero_action": "25SF1353",
        "date_redaction": datetime.now().strftime('%d/%m/%Y'),
        "observations_groupe": "Groupe test - Automatisation réussie !",
        "themes_modules": "CP1 - Réception de marchandises\nContrôle quantitatif et qualitatif\nUtilisation du scan",
        "previsions": "CP2 - Stockage et rangement en zone dédiée",
        "apprenants": [
            {
                "nom": "DUPONT",
                "prenom": "Jean",
                "observation": "Très investi, bonne compréhension des concepts"
            },
            {
                "nom": "MARTIN",
                "prenom": "Sophie",
                "observation": "Quelques difficultés sur le scan, à accompagner"
            },
            {
                "nom": "BERNARD",
                "prenom": "Lucas",
                "observation": "Absent cette semaine"
            }
        ]
    }
    
    try:
        output_path, output_filename = fill_suivi_formation(test_data, TEMPLATE_PATH)
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({
            'error': str(e),
            'type': type(e).__name__
        }), 500


if __name__ == '__main__':
    # Vérifier que le template existe
    if not os.path.exists(TEMPLATE_PATH):
        print(f"⚠️  ATTENTION: Le fichier template n'existe pas: {TEMPLATE_PATH}")
        print("Veuillez définir la variable d'environnement TEMPLATE_PATH")
    else:
        print(f"✅ Template trouvé: {TEMPLATE_PATH}")
    
    # Démarrer le serveur
    port = int(os.environ.get('PORT', 5000))
    print(f"🚀 Serveur démarré sur le port {port}")
    app.run(host='0.0.0.0', port=port, debug=True)
